# invoice_generator.py - Red headings, letterhead in PDF, keep stamp with total

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd
import numpy as np
from io import BytesIO
import os
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak, Image, KeepTogether
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT


def get_template_path():
    """Get the path to the invoice template"""
    current_dir = os.path.dirname(os.path.abspath(__file__))
    template_path = os.path.join(current_dir, 'templates', 'invoice_template.docx')

    if os.path.exists(template_path):
        return template_path
    else:
        return None


def get_stamp_path():
    """Get the path to the stamp image"""
    current_dir = os.path.dirname(os.path.abspath(__file__))
    stamp_path = os.path.join(current_dir, 'templates', 'stamp.png')

    if os.path.exists(stamp_path):
        return stamp_path
    else:
        return None


def safe_int(value):
    """Safely convert value to int, return 0 if NaN"""
    try:
        if pd.isna(value):
            return 0
        return int(value)
    except:
        return 0


def safe_float(value):
    """Safely convert value to float, return 0.0 if NaN"""
    try:
        if pd.isna(value):
            return 0.0
        return float(value)
    except:
        return 0.0


def safe_str(value):
    """Safely convert value to string, return empty string if NaN"""
    try:
        if pd.isna(value):
            return ''
        return str(value)
    except:
        return ''


def number_to_words(number):
    """Convert number to words for amount - whole numbers only"""
    ones = ['', 'ONE', 'TWO', 'THREE', 'FOUR', 'FIVE', 'SIX', 'SEVEN', 'EIGHT', 'NINE']
    teens = ['TEN', 'ELEVEN', 'TWELVE', 'THIRTEEN', 'FOURTEEN', 'FIFTEEN', 'SIXTEEN', 'SEVENTEEN', 'EIGHTEEN',
             'NINETEEN']
    tens = ['', '', 'TWENTY', 'THIRTY', 'FORTY', 'FIFTY', 'SIXTY', 'SEVENTY', 'EIGHTY', 'NINETY']

    def convert_below_thousand(n):
        if n == 0:
            return ''
        elif n < 10:
            return ones[n]
        elif n < 20:
            return teens[n - 10]
        elif n < 100:
            return tens[n // 10] + (' ' + ones[n % 10] if n % 10 != 0 else '')
        else:
            return ones[n // 100] + ' HUNDRED' + (' AND ' + convert_below_thousand(n % 100) if n % 100 != 0 else '')

    # Round to nearest whole number
    number = round(number)

    if number == 0:
        return 'ZERO ONLY'

    result = []

    if number >= 1000000:
        millions = number // 1000000
        result.append(convert_below_thousand(millions) + ' MILLION')
        number %= 1000000

    if number >= 1000:
        thousands = number // 1000
        result.append(convert_below_thousand(thousands) + ' THOUSAND')
        number %= 1000

    if number > 0:
        result.append(convert_below_thousand(number))

    return ' '.join(result) + ' ONLY'


def generate_combined_word(df, invoice_data):
    """
    Generate Word document with both Invoice and Delivery Note using letterhead template
    """
    df = df.copy()

    # Clean data
    df = df.dropna(subset=['Item No.', 'Item Description'])
    df['Quantity'] = df['Quantity'].fillna(0)
    df['Unit Price'] = df['Unit Price'].fillna(0)

    # Get template
    template_path = get_template_path()
    stamp_path = get_stamp_path()

    if template_path and os.path.exists(template_path):
        doc = Document(template_path)
    else:
        raise FileNotFoundError("Template file 'invoice_template.docx' not found in 'templates' folder!")

    # ========================================
    # INVOICE SECTION
    # ========================================

    doc.add_paragraph()

    # Invoice header - CENTERED and RED
    inv_header = doc.add_paragraph()
    run = inv_header.add_run(f'Invoice no: {invoice_data["invoice_number"]}')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(255, 0, 0)  # RED COLOR
    inv_header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Recipient and details
    doc.add_paragraph(f'To: Master / manager – {invoice_data["vessel_name"]}')
    doc.add_paragraph(f'{invoice_data.get("invoice_type", "Ship provision")}')
    doc.add_paragraph(f'Port of delivery: {invoice_data["port_of_delivery"]}')

    doc.add_paragraph()

    # Calculate totals
    df['Total'] = df['Quantity'] * df['Unit Price']
    grand_total = df['Total'].sum()

    # Add serial numbers
    if 'Serial No.' not in df.columns:
        df.insert(0, 'Serial No.', range(1, len(df) + 1))

    # Create invoice table
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'

    # Header row
    headers = ['Serial\nNo.', 'Item No.', 'Item Description', 'Quantity', 'UoM\nCode', 'Unit\nPrice',
               f'Total\n({invoice_data["currency"]})']
    header_cells = table.rows[0].cells

    for i, header in enumerate(headers):
        cell = header_cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)

    # Data rows
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(safe_int(row['Serial No.']))
        row_cells[1].text = safe_str(row['Item No.'])
        row_cells[2].text = safe_str(row['Item Description'])
        row_cells[3].text = str(safe_int(row['Quantity']))
        row_cells[4].text = safe_str(row['UoM Code'])
        row_cells[5].text = f"{safe_float(row['Unit Price']):.2f}"
        row_cells[6].text = f"{safe_float(row['Total']):.2f}"

        for i, cell in enumerate(row_cells):
            for paragraph in cell.paragraphs:
                if i == 2:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()

    # Total
    total_para = doc.add_paragraph()
    run = total_para.add_run(f'TOTAL {invoice_data["currency"]} {grand_total:.2f}')
    run.bold = True
    run.font.size = Pt(11)
    total_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Amount in words - BOLD
    currency_name = "US DOLLARS" if invoice_data["currency"] == "USD" else "UAE DIRHAMS"
    amount_words = number_to_words(grand_total)
    words_para = doc.add_paragraph()
    run = words_para.add_run(f'AMOUNT IN WORDS: {currency_name} {amount_words}')
    run.bold = True
    run.font.size = Pt(10)

    # Add stamp for invoice - keep on same page
    if stamp_path:
        stamp_para = doc.add_paragraph()
        stamp_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = stamp_para.add_run()
        run.add_picture(stamp_path, width=Inches(1.5))
        # Set paragraph to keep with previous to prevent page break
        stamp_para.paragraph_format.keep_with_next = True

    # ========================================
    # PAGE BREAK
    # ========================================
    doc.add_page_break()

    # ========================================
    # DELIVERY NOTE SECTION
    # ========================================

    doc.add_paragraph()

    # Delivery note header - CENTERED and RED
    dn_header = doc.add_paragraph()
    run = dn_header.add_run(f'Delivery note: {invoice_data["invoice_number"]}-A')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(255, 0, 0)  # RED COLOR
    dn_header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Recipient and details
    doc.add_paragraph(f'To: Master / manager – {invoice_data["vessel_name"]}')
    doc.add_paragraph(f'{invoice_data.get("invoice_type", "Ship provision")}')
    doc.add_paragraph(f'Port of delivery: {invoice_data["port_of_delivery"]}')

    doc.add_paragraph()

    # Create delivery note table - NO UoM column
    dn_table = doc.add_table(rows=1, cols=4)
    dn_table.style = 'Table Grid'

    # Header row
    dn_headers = ['Serial No.', 'Item No.', 'Item Description', 'Quantity']
    dn_header_cells = dn_table.rows[0].cells

    for i, header in enumerate(dn_headers):
        cell = dn_header_cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)

    # Data rows
    for _, row in df.iterrows():
        row_cells = dn_table.add_row().cells
        row_cells[0].text = str(safe_int(row['Serial No.']))
        row_cells[1].text = safe_str(row['Item No.'])
        row_cells[2].text = safe_str(row['Item Description'])
        row_cells[3].text = str(safe_int(row['Quantity']))

        for i, cell in enumerate(row_cells):
            for paragraph in cell.paragraphs:
                if i == 2:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    # Add stamp for delivery note - keep on same page
    if stamp_path:
        doc.add_paragraph()
        stamp_para = doc.add_paragraph()
        stamp_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = stamp_para.add_run()
        run.add_picture(stamp_path, width=Inches(1.5))

    # Save to BytesIO
    output = BytesIO()
    doc.save(output)
    output.seek(0)

    return output, grand_total


def generate_combined_pdf(df, invoice_data):
    """
    Generate PDF by first creating Word document, then converting to PDF (no dialogs)
    """
    import tempfile
    import shutil
    import subprocess
    import platform

    # Create a temporary directory for file operations
    temp_dir = tempfile.mkdtemp()

    try:
        # Generate Word document first (this already has the letterhead template!)
        word_output, grand_total = generate_combined_word(df, invoice_data)

        # Save Word document to temporary file
        temp_docx_path = os.path.join(temp_dir, 'invoice_temp.docx')
        with open(temp_docx_path, 'wb') as f:
            f.write(word_output.getvalue())

        # Convert to PDF
        temp_pdf_path = os.path.join(temp_dir, 'invoice_temp.pdf')

        conversion_success = False

        # Method 1: LibreOffice (BEST - truly headless, no dialogs)
        try:
            # Different possible LibreOffice paths
            libreoffice_cmds = []

            if platform.system() == 'Windows':
                libreoffice_cmds = [
                    'C:\\Program Files\\LibreOffice\\program\\soffice.exe',
                    'C:\\Program Files (x86)\\LibreOffice\\program\\soffice.exe',
                    'soffice.exe',
                    'libreoffice.exe'
                ]
            elif platform.system() == 'Darwin':  # Mac
                libreoffice_cmds = [
                    '/Applications/LibreOffice.app/Contents/MacOS/soffice',
                ]
            else:  # Linux
                libreoffice_cmds = [
                    'libreoffice',
                    'soffice',
                    '/usr/bin/libreoffice',
                    '/usr/bin/soffice'
                ]

            for cmd in libreoffice_cmds:
                try:
                    # Use --invisible flag to prevent any UI from showing
                    result = subprocess.run([
                        cmd,
                        '--headless',
                        '--invisible',
                        '--nodefault',
                        '--nofirststartwizard',
                        '--nolockcheck',
                        '--nologo',
                        '--norestore',
                        '--convert-to', 'pdf',
                        '--outdir', temp_dir,
                        temp_docx_path
                    ],
                        check=True,
                        capture_output=True,
                        timeout=30,
                        creationflags=subprocess.CREATE_NO_WINDOW if platform.system() == 'Windows' else 0
                    )
                    conversion_success = True
                    print(f"Converted using LibreOffice (headless): {cmd}")
                    break
                except (subprocess.CalledProcessError, FileNotFoundError, subprocess.TimeoutExpired) as e:
                    continue

        except Exception as e:
            print(f"LibreOffice conversion failed: {e}")

        # Method 2: comtypes (Windows only - uses Word automation without showing UI)
        if not conversion_success and platform.system() == 'Windows':
            try:
                import comtypes.client

                # Initialize Word application
                word = comtypes.client.CreateObject('Word.Application')
                word.Visible = False  # Don't show Word window
                word.DisplayAlerts = 0  # Don't show any alerts/dialogs

                # Open document
                doc = word.Documents.Open(temp_docx_path)

                # Save as PDF (17 is the PDF format constant)
                doc.SaveAs(temp_pdf_path, FileFormat=17)

                # Close document and quit Word
                doc.Close()
                word.Quit()

                conversion_success = True
                print("Converted using Word COM automation")

            except Exception as e:
                print(f"Word COM automation failed: {e}")

        # Method 3: docx2pdf with Word (Windows - may show dialog)
        if not conversion_success and platform.system() == 'Windows':
            try:
                from docx2pdf import convert
                convert(temp_docx_path, temp_pdf_path)
                conversion_success = True
                print("Converted using docx2pdf")
            except Exception as e:
                print(f"docx2pdf failed: {e}")

        if not conversion_success:
            raise Exception(
                "PDF conversion failed. Please install LibreOffice:\n"
                "- Windows: Download from https://www.libreoffice.org/download/\n"
                "- Linux: sudo apt-get install libreoffice\n"
                "- Mac: brew install --cask libreoffice\n"
                "\nFor Windows, you can also: pip install comtypes"
            )

        # Read the generated PDF
        if not os.path.exists(temp_pdf_path):
            raise FileNotFoundError(f"PDF file was not created at {temp_pdf_path}")

        with open(temp_pdf_path, 'rb') as f:
            pdf_output = BytesIO(f.read())

        pdf_output.seek(0)
        return pdf_output, grand_total

    finally:
        # Clean up temporary directory
        try:
            shutil.rmtree(temp_dir)
        except Exception as e:
            print(f"Could not clean up temp directory: {e}")